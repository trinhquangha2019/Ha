# Script xuất nội dung ra file Word (.docx)
# Tác giả: Dasi

import subprocess
import sys

# Cài đặt thư viện nếu chưa có
def install_package(package):
    subprocess.check_call([sys.executable, "-m", "pip", "install", package, "-q"])

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Đang cài đặt python-docx...")
    install_package("python-docx")
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_docx():
    # Tạo document mới
    doc = Document()
    
    # ===== TIÊU ĐỀ =====
    title = doc.add_heading('ĐỀ XUẤT CONTENT WEBSITE NHÀ HÀNG RẠN BIỂN 2025', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Thông tin
    doc.add_paragraph('Tác giả: Dasi')
    doc.add_paragraph('Ngày tạo: 04/12/2025')
    doc.add_paragraph('Nguồn tham khảo: haisanranbien.vn')
    doc.add_paragraph()
    
    # ===== NHÓM 1 =====
    doc.add_heading('NHÓM 1: KIẾN THỨC HẢI SẢN (Giáo dục khách hàng)', level=1)
    
    table1 = doc.add_table(rows=6, cols=3)
    table1.style = 'Table Grid'
    
    # Header
    table1.rows[0].cells[0].text = '#'
    table1.rows[0].cells[1].text = 'Chủ đề'
    table1.rows[0].cells[2].text = 'Mục đích'
    
    # Data
    data1 = [
        ('1', 'Cách chọn hải sản tươi sống', 'Xây dựng trust, khách tin tưởng'),
        ('2', 'Mùa vụ hải sản Việt Nam', 'Khách biết khi nào ăn gì ngon nhất'),
        ('3', 'Giá trị dinh dưỡng từng loại', 'Thu hút khách quan tâm sức khỏe'),
        ('4', 'Phân biệt hải sản thật - giả', 'Khẳng định uy tín nhà hàng'),
        ('5', 'Xuất xứ các vùng biển VN', 'Giáo dục về Trường Sa, Phú Quý...'),
    ]
    for i, row_data in enumerate(data1):
        table1.rows[i+1].cells[0].text = row_data[0]
        table1.rows[i+1].cells[1].text = row_data[1]
        table1.rows[i+1].cells[2].text = row_data[2]
    
    doc.add_paragraph()
    
    # ===== NHÓM 2 =====
    doc.add_heading('NHÓM 2: CÔNG THỨC & CHẾ BIẾN', level=1)
    
    table2 = doc.add_table(rows=6, cols=3)
    table2.style = 'Table Grid'
    
    table2.rows[0].cells[0].text = '#'
    table2.rows[0].cells[1].text = 'Chủ đề'
    table2.rows[0].cells[2].text = 'Mục đích'
    
    data2 = [
        ('1', 'Nghệ thuật Sashimi', 'Highlight đầu bếp chuyên nghiệp'),
        ('2', 'Các cách chế biến phổ biến', 'Hấp, nướng, xào, chiên...'),
        ('3', 'Nước chấm đặc biệt Rạn Biển', 'Tạo sự khác biệt'),
        ('4', 'Cách ăn hải sản đúng cách', 'Giáo dục thực khách'),
        ('5', 'Set menu gợi ý theo số người', 'Hỗ trợ đặt bàn'),
    ]
    for i, row_data in enumerate(data2):
        table2.rows[i+1].cells[0].text = row_data[0]
        table2.rows[i+1].cells[1].text = row_data[1]
        table2.rows[i+1].cells[2].text = row_data[2]
    
    doc.add_paragraph()
    
    # ===== NHÓM 3 =====
    doc.add_heading('NHÓM 3: DỊCH VỤ & SỰ KIỆN', level=1)
    
    table3 = doc.add_table(rows=7, cols=3)
    table3.style = 'Table Grid'
    
    table3.rows[0].cells[0].text = '#'
    table3.rows[0].cells[1].text = 'Chủ đề'
    table3.rows[0].cells[2].text = 'Mục đích'
    
    data3 = [
        ('1', 'Tổ chức tiệc công ty', 'Thu hút doanh nghiệp'),
        ('2', 'Tiệc sinh nhật / kỷ niệm', 'Thu hút gia đình'),
        ('3', 'Đặt tiệc cưới hỏi', 'Mở rộng dịch vụ'),
        ('4', 'Phòng VIP & sức chứa', 'Thông tin chi tiết từng chi nhánh'),
        ('5', 'Dịch vụ giao hàng tận nơi', 'Nếu có'),
        ('6', 'Voucher / Quà tặng', 'Tăng doanh thu'),
    ]
    for i, row_data in enumerate(data3):
        table3.rows[i+1].cells[0].text = row_data[0]
        table3.rows[i+1].cells[1].text = row_data[1]
        table3.rows[i+1].cells[2].text = row_data[2]
    
    doc.add_paragraph()
    
    # ===== NHÓM 4 =====
    doc.add_heading('NHÓM 4: CÂU CHUYỆN THƯƠNG HIỆU', level=1)
    
    table4 = doc.add_table(rows=6, cols=3)
    table4.style = 'Table Grid'
    
    table4.rows[0].cells[0].text = '#'
    table4.rows[0].cells[1].text = 'Chủ đề'
    table4.rows[0].cells[2].text = 'Mục đích'
    
    data4 = [
        ('1', 'Câu chuyện Rạn Biển', 'Brand storytelling'),
        ('2', 'Đội ngũ đầu bếp', 'Giới thiệu chef, tạo kết nối'),
        ('3', 'Hành trình từ biển đến bàn ăn', 'Quy trình đảm bảo tươi sống'),
        ('4', 'Cam kết chất lượng', 'Trust building'),
        ('5', 'Đối tác cung cấp hải sản', 'Nguồn gốc rõ ràng'),
    ]
    for i, row_data in enumerate(data4):
        table4.rows[i+1].cells[0].text = row_data[0]
        table4.rows[i+1].cells[1].text = row_data[1]
        table4.rows[i+1].cells[2].text = row_data[2]
    
    doc.add_paragraph()
    
    # ===== NHÓM 5 =====
    doc.add_heading('NHÓM 5: THÔNG TIN CHI NHÁNH', level=1)
    
    table5 = doc.add_table(rows=6, cols=3)
    table5.style = 'Table Grid'
    
    table5.rows[0].cells[0].text = '#'
    table5.rows[0].cells[1].text = 'Chủ đề'
    table5.rows[0].cells[2].text = 'Mục đích'
    
    data5 = [
        ('1', 'Giới thiệu từng chi nhánh', 'Ảnh, video, không gian'),
        ('2', 'Bản đồ & hướng dẫn đi', 'SEO địa phương'),
        ('3', 'Sức chứa & phòng VIP', 'Hỗ trợ đặt bàn'),
        ('4', 'Bãi đỗ xe', 'Thông tin thiết thực'),
        ('5', 'Review từng chi nhánh', 'Social proof'),
    ]
    for i, row_data in enumerate(data5):
        table5.rows[i+1].cells[0].text = row_data[0]
        table5.rows[i+1].cells[1].text = row_data[1]
        table5.rows[i+1].cells[2].text = row_data[2]
    
    doc.add_paragraph()
    
    # ===== NHÓM 6 =====
    doc.add_heading('NHÓM 6: FAQ & HỖ TRỢ', level=1)
    
    table6 = doc.add_table(rows=7, cols=3)
    table6.style = 'Table Grid'
    
    table6.rows[0].cells[0].text = '#'
    table6.rows[0].cells[1].text = 'Câu hỏi thường gặp'
    table6.rows[0].cells[2].text = 'Mục đích'
    
    data6 = [
        ('1', 'Giá hải sản tính như thế nào?', 'Giải đáp thắc mắc'),
        ('2', 'Có cần đặt bàn trước không?', 'Hướng dẫn khách'),
        ('3', 'Nhà hàng có ship không?', 'Dịch vụ'),
        ('4', 'Thanh toán những hình thức nào?', 'Tiện lợi'),
        ('5', 'Có menu cho trẻ em không?', 'Thu hút gia đình'),
        ('6', 'Có món chay/ăn kiêng không?', 'Đa dạng khách hàng'),
    ]
    for i, row_data in enumerate(data6):
        table6.rows[i+1].cells[0].text = row_data[0]
        table6.rows[i+1].cells[1].text = row_data[1]
        table6.rows[i+1].cells[2].text = row_data[2]
    
    doc.add_paragraph()
    
    # ===== NHÓM 7 =====
    doc.add_heading('NHÓM 7: NỘI DUNG ĐA PHƯƠNG TIỆN', level=1)
    
    table7 = doc.add_table(rows=7, cols=3)
    table7.style = 'Table Grid'
    
    table7.rows[0].cells[0].text = '#'
    table7.rows[0].cells[1].text = 'Loại'
    table7.rows[0].cells[2].text = 'Ý tưởng'
    
    data7 = [
        ('1', 'Video', 'Đầu bếp làm Sashimi tại bàn'),
        ('2', 'Video', 'Tour nhà hàng 360°'),
        ('3', 'Video', 'Hải sản bơi trong hồ'),
        ('4', 'Gallery', 'Món ăn đẹp mắt'),
        ('5', 'Gallery', 'Không gian phòng VIP'),
        ('6', 'Infographic', 'So sánh các loại tôm hùm'),
    ]
    for i, row_data in enumerate(data7):
        table7.rows[i+1].cells[0].text = row_data[0]
        table7.rows[i+1].cells[1].text = row_data[1]
        table7.rows[i+1].cells[2].text = row_data[2]
    
    doc.add_paragraph()
    
    # ===== NHÓM 8 =====
    doc.add_heading('NHÓM 8: NỘI DUNG THEO MÙA', level=1)
    
    table8 = doc.add_table(rows=7, cols=3)
    table8.style = 'Table Grid'
    
    table8.rows[0].cells[0].text = 'Tháng'
    table8.rows[0].cells[1].text = 'Chủ đề'
    table8.rows[0].cells[2].text = 'Ghi chú'
    
    data8 = [
        ('Tết', 'Set menu Tết, tiệc tất niên', 'Quan trọng nhất'),
        ('Valentine', 'Bữa tối lãng mạn cho đôi', 'Cặp đôi'),
        ('8/3 - 20/10', 'Tri ân phụ nữ', 'Gia đình'),
        ('Hè', 'Hải sản mùa hè, du khách', 'Du lịch'),
        ('Giáng Sinh', 'Tiệc Noel, năm mới', 'Cuối năm'),
        ('Trung Thu', 'Tiệc gia đình', 'Truyền thống'),
    ]
    for i, row_data in enumerate(data8):
        table8.rows[i+1].cells[0].text = row_data[0]
        table8.rows[i+1].cells[1].text = row_data[1]
        table8.rows[i+1].cells[2].text = row_data[2]
    
    doc.add_paragraph()
    
    # ===== NHÓM 9 =====
    doc.add_heading('NHÓM 9: NỘI DUNG ĐA NGÔN NGỮ', level=1)
    doc.add_paragraph('Website đã có: Tiếng Việt, English, Chinese, Korean')
    
    table9 = doc.add_table(rows=6, cols=3)
    table9.style = 'Table Grid'
    
    table9.rows[0].cells[0].text = '#'
    table9.rows[0].cells[1].text = 'Nội dung cần dịch'
    table9.rows[0].cells[2].text = 'Ưu tiên'
    
    data9 = [
        ('1', 'Thực đơn đầy đủ', '⭐⭐⭐'),
        ('2', 'Giới thiệu nhà hàng', '⭐⭐⭐'),
        ('3', 'Bài viết hải sản phổ biến', '⭐⭐'),
        ('4', 'Hướng dẫn đặt bàn', '⭐⭐⭐'),
        ('5', 'FAQ', '⭐⭐'),
    ]
    for i, row_data in enumerate(data9):
        table9.rows[i+1].cells[0].text = row_data[0]
        table9.rows[i+1].cells[1].text = row_data[1]
        table9.rows[i+1].cells[2].text = row_data[2]
    
    doc.add_paragraph()
    
    # ===== TỔNG HỢP ƯU TIÊN =====
    doc.add_heading('TỔNG HỢP ƯU TIÊN', level=1)
    
    doc.add_heading('CẦN LÀM NGAY (Tăng chuyển đổi)', level=2)
    doc.add_paragraph('1. Trang đặt tiệc công ty - Doanh thu lớn')
    doc.add_paragraph('2. FAQ đầy đủ - Giảm hỏi đáp, tăng UX')
    doc.add_paragraph('3. Giới thiệu từng chi nhánh - SEO địa phương')
    doc.add_paragraph('4. Set menu gợi ý - Hỗ trợ quyết định')
    
    doc.add_heading('NÊN LÀM (Xây dựng thương hiệu)', level=2)
    doc.add_paragraph('1. Câu chuyện thương hiệu - Brand building')
    doc.add_paragraph('2. Giới thiệu đầu bếp - Tạo kết nối')
    doc.add_paragraph('3. Video sashimi - Viral content')
    doc.add_paragraph('4. Mùa vụ hải sản - SEO + Giáo dục')
    
    doc.add_heading('CÓ THỂ LÀM SAU (Bổ sung)', level=2)
    doc.add_paragraph('1. Blog công thức - Long-term SEO')
    doc.add_paragraph('2. Infographic - Social sharing')
    doc.add_paragraph('3. Nội dung đa ngôn ngữ - Du khách')
    
    doc.add_paragraph()
    
    # ===== KẾ HOẠCH TRIỂN KHAI =====
    doc.add_heading('KẾ HOẠCH TRIỂN KHAI 2025', level=1)
    
    doc.add_heading('Quý 1/2025 (Tháng 1-3)', level=2)
    doc.add_paragraph('☐ Trang đặt tiệc công ty')
    doc.add_paragraph('☐ FAQ đầy đủ')
    doc.add_paragraph('☐ Set menu Tết')
    doc.add_paragraph('☐ Giới thiệu chi nhánh Trung Tâm')
    
    doc.add_heading('Quý 2/2025 (Tháng 4-6)', level=2)
    doc.add_paragraph('☐ Câu chuyện thương hiệu')
    doc.add_paragraph('☐ Giới thiệu đầu bếp')
    doc.add_paragraph('☐ Kiến thức hải sản (5 bài)')
    doc.add_paragraph('☐ Giới thiệu các chi nhánh còn lại')
    
    doc.add_heading('Quý 3/2025 (Tháng 7-9)', level=2)
    doc.add_paragraph('☐ Video sashimi')
    doc.add_paragraph('☐ Mùa vụ hải sản')
    doc.add_paragraph('☐ Gallery món ăn')
    doc.add_paragraph('☐ Nội dung hè/du lịch')
    
    doc.add_heading('Quý 4/2025 (Tháng 10-12)', level=2)
    doc.add_paragraph('☐ Nội dung Giáng Sinh/Năm mới')
    doc.add_paragraph('☐ Infographic')
    doc.add_paragraph('☐ Dịch nội dung đa ngôn ngữ')
    doc.add_paragraph('☐ Tổng kết & đánh giá')
    
    # Lưu file
    output_path = r'C:\AI\De_xuat_content_website_2025.docx'
    doc.save(output_path)
    print(f"✅ Đã xuất file thành công!")
    print(f"📄 File: {output_path}")
    return output_path

if __name__ == "__main__":
    create_docx()



