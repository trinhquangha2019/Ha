"""
Tool đọc file Word (.docx) và extract text
Có thể export lên NotebookLM
"""

import os
import sys
from pathlib import Path

def read_docx(docx_path):
    """Đọc file Word (.docx) và extract text"""
    try:
        from docx import Document
        
        print(f"📄 Đang đọc file: {docx_path}")
        doc = Document(docx_path)
        
        text_content = []
        total_paragraphs = len(doc.paragraphs)
        print(f"📑 Tổng số đoạn văn: {total_paragraphs}")
        
        for i, para in enumerate(doc.paragraphs, 1):
            text = para.text.strip()
            if text:
                text_content.append(text)
                if i % 50 == 0:
                    print(f"   Đã đọc {i}/{total_paragraphs} đoạn...", end='\r')
        
        # Đọc text từ tables nếu có
        if doc.tables:
            print(f"\n📊 Tìm thấy {len(doc.tables)} bảng, đang đọc...")
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_content.append(" | ".join(row_text))
        
        result = "\n".join(text_content)
        
        if not result:
            print("⚠️  File không có nội dung text")
            return None
        
        print(f"\n✅ Đã đọc {len(result)} ký tự từ file Word")
        return result
        
    except ImportError:
        print("⚠️  Đang cài đặt python-docx...")
        os.system("pip install python-docx")
        return read_docx(docx_path)
    except Exception as e:
        print(f"❌ Lỗi khi đọc file Word: {e}")
        return None

def main():
    """Hàm chính"""
    print("=" * 60)
    print("TOOL ĐỌC FILE WORD (.DOCX)")
    print("=" * 60)
    
    # Nhận đường dẫn file
    if len(sys.argv) > 1:
        docx_path = sys.argv[1]
    else:
        docx_path = input("\nNhập đường dẫn file .docx: ").strip().strip('"')
    
    if not os.path.exists(docx_path):
        print(f"❌ Không tìm thấy file: {docx_path}")
        return
    
    if not docx_path.lower().endswith('.docx'):
        print(f"⚠️  File không phải .docx, nhưng sẽ thử đọc...")
    
    text_content = read_docx(docx_path)
    
    if not text_content:
        print("❌ Không thể đọc được nội dung từ file")
        return
    
    print(f"\n✅ Đã đọc {len(text_content)} ký tự\n")
    
    # Menu lựa chọn
    print("Chọn hành động:")
    print("1. Chỉ lưu text ra file")
    print("2. Export lên NotebookLM")
    print("3. Cả hai")
    
    choice = None
    try:
        choice = input("\nLựa chọn (1/2/3, mặc định 1): ").strip()
    except (EOFError, KeyboardInterrupt):
        print("\n⚠️  Không có input, tự động lưu text ra file...")
        choice = '1'
    
    if not choice or choice not in ['1', '2', '3']:
        choice = '1'
    
    docx_name = Path(docx_path).stem
    output_text = f"{docx_name}_extracted.txt"
    
    # Lưu text
    if choice in ['1', '3']:
        try:
            output_path = os.path.abspath(output_text)
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text_content)
            print(f"✅ Đã lưu text vào: {output_path}")
            print(f"   ({len(text_content)} ký tự)")
        except Exception as e:
            print(f"❌ Lỗi khi lưu file: {e}")
    
    # Export lên NotebookLM
    if choice in ['2', '3']:
        try:
            from notebooklm_integration import export_to_notebooklm, show_notebooklm_instructions
            
            print("\n" + "=" * 60)
            print("📚 EXPORT LÊN NOTEBOOKLM")
            print("=" * 60)
            
            print("\nChọn phương thức export:")
            print("1. Tạo file text (copy/paste vào NotebookLM)")
            print("2. Upload lên Google Drive (cần setup API)")
            print("3. Cả hai")
            
            export_method = input("\nLựa chọn (1/2/3, mặc định 1): ").strip()
            if not export_method or export_method not in ['1', '2', '3']:
                export_method = '1'
            
            method_map = {
                '1': 'file',
                '2': 'drive',
                '3': 'both'
            }
            
            method = method_map[export_method]
            output_notebooklm = f"{docx_name}_notebooklm.txt"
            
            result = export_to_notebooklm(
                text_content,
                method=method,
                pdf_path=docx_path,  # Dùng docx_path cho source name
                output_file=output_notebooklm
            )
            
            if result['success']:
                print("\n✅ Export thành công!")
                if result['file_path']:
                    print(f"📄 File: {result['file_path']}")
                if result['drive_link']:
                    print(f"🔗 Google Drive: {result['drive_link']}")
                
                print("\n💡 Bước tiếp theo:")
                show_notebooklm_instructions()
            else:
                print("\n⚠️  Export không thành công. Vui lòng thử lại.")
                
        except ImportError:
            print("\n⚠️  Module notebooklm_integration không tìm thấy.")
            print("   Tạo file text thủ công...")
            output_notebooklm = f"{docx_name}_notebooklm.txt"
            try:
                with open(output_notebooklm, 'w', encoding='utf-8') as f:
                    f.write(text_content)
                print(f"✅ Đã tạo file: {os.path.abspath(output_notebooklm)}")
                print("   Bạn có thể copy nội dung và paste vào NotebookLM")
            except Exception as e:
                print(f"❌ Lỗi: {e}")
        except Exception as e:
            print(f"\n❌ Lỗi khi export lên NotebookLM: {e}")

if __name__ == "__main__":
    main()



